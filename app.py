import streamlit as st
from crewai import Agent, Task, Crew, Process
import os
from crewai_tools import ScrapeWebsiteTool, SerperDevTool
from dotenv import load_dotenv
from langchain_openai import ChatOpenAI
from docx import Document
from io import BytesIO
import base64

# Load environment variables
load_dotenv()

# Configure API keys
os.environ["OPENAI_API_KEY"] = os.getenv("OPENAI_API_KEY")
os.environ["SERPER_API_KEY"] = os.getenv("SERPER_API_KEY")

# Helper Functions
def generate_docx(result):
    doc = Document()
    doc.add_heading('Healthcare Diagnosis and Treatment Recommendations', 0)
    
    # Convert result to string if it's a tuple
    if isinstance(result, tuple):
        result_str = '\n\n'.join(str(item) for item in result)
    else:
        result_str = str(result)
    
    doc.add_paragraph(result_str)
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

def get_download_link(bio, filename):
    b64 = base64.b64encode(bio.read()).decode()
    return f'<a href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{b64}" download="{filename}" class="download-button">Download Report</a>'

# Page Configuration
st.set_page_config(
    page_title="Medical AI Assistant",
    page_icon="🏥",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Custom CSS
st.markdown("""
    <style>
    .main {
        padding: 2rem;
    }
    .stButton > button {
        width: 100%;
        background-color: #007bff;
        color: white;
        padding: 0.5rem 1rem;
        border-radius: 0.5rem;
        border: none;
        margin-top: 1rem;
    }
    .stButton > button:hover {
        background-color: #0056b3;
    }
    .download-button {
        display: inline-block;
        padding: 0.5rem 1rem;
        background-color: #28a745;
        color: white;
        text-decoration: none;
        border-radius: 0.5rem;
        margin-top: 1rem;
    }
    .download-button:hover {
        background-color: #218838;
        color: white;
    }
    .stTextInput > div > div > input {
        border-radius: 0.5rem;
    }
    .stTextArea > div > div > textarea {
        border-radius: 0.5rem;
    }
    </style>
""", unsafe_allow_html=True)

# Sidebar for Patient Information
with st.sidebar:
    st.image("https://formaspace.com/wp-content/uploads/2024/04/ai-dr.jpeg", width=100)
    st.title("Patient Information")
    
    with st.form("patient_info"):
        gender = st.selectbox('Gender', ('Male', 'Female', 'Other'))
        age = st.number_input('Age', min_value=0, max_value=120, value=25)
        height = st.number_input('Height (cm)', min_value=0, max_value=300, value=170)
        weight = st.number_input('Weight (kg)', min_value=0, max_value=500, value=70)
        submit_button = st.form_submit_button("Save Patient Info")

# Main Content
st.title("🏥 Medical AI Assistant")
st.markdown("""
    <div style='background-color: #f8f9fa; padding: 1rem; border-radius: 0.5rem; margin-bottom: 2rem;'>
        <h4>Welcome to the Medical AI Assistant</h4>
        <p>This AI-powered system helps medical professionals with diagnosis and treatment recommendations. 
        Please enter the patient's symptoms and medical history below.</p>
    </div>
""", unsafe_allow_html=True)

# Create tabs for different sections
tab1, tab2 = st.tabs(["📝 Patient Assessment", "📊 Results"])

with tab1:
    col1, col2 = st.columns(2)
    
    with col1:
        st.subheader("Current Symptoms")
        symptoms = st.text_area(
            'Describe the symptoms in detail',
            placeholder='e.g., persistent fever for 3 days, dry cough, fatigue',
            height=200
        )

    with col2:
        st.subheader("Medical History")
        medical_history = st.text_area(
            'Enter relevant medical history',
            placeholder='e.g., Type 2 diabetes diagnosed in 2019, hypertension',
            height=200
        )

    # Additional Information
    with st.expander("Additional Information (Optional)"):
        col3, col4 = st.columns(2)
        with col3:
            allergies = st.text_area("Known Allergies", placeholder="e.g., penicillin, peanuts")
            current_medications = st.text_area("Current Medications", placeholder="e.g., metformin 500mg twice daily")
        with col4:
            family_history = st.text_area("Family History", placeholder="e.g., heart disease, diabetes")
            lifestyle = st.text_area("Lifestyle Factors", placeholder="e.g., smoker, exercises 3 times a week")

# Initialize Tools and Agents
search_tool = SerperDevTool()
scrape_tool = ScrapeWebsiteTool()

llm = ChatOpenAI(
    model="gpt-3.5-turbo-16k",
    temperature=0.1,
    max_tokens=8000
)

# Define Agents
diagnostician = Agent(
    role="Medical Diagnostician",
    goal="Analyze patient symptoms and medical history to provide a preliminary diagnosis.",
    backstory="Expert in diagnosing medical conditions using advanced algorithms and comprehensive medical knowledge.",
    verbose=True,
    allow_delegation=False,
    tools=[search_tool, scrape_tool],
    llm=llm
)

treatment_advisor = Agent(
    role="Treatment Advisor",
    goal="Recommend appropriate treatment plans based on the diagnosis.",
    backstory="Specialist in creating personalized treatment plans considering patient history and current medical best practices.",
    verbose=True,
    allow_delegation=False,
    tools=[search_tool, scrape_tool],
    llm=llm
)

# Define Tasks
diagnose_task = Task(
    description=(
        f"1. Analyze the patient's symptoms ({symptoms}) and medical history ({medical_history}).\n"
        f"2. Consider additional factors: Age: {age}, Gender: {gender}\n"
        "3. Provide a preliminary diagnosis with possible conditions.\n"
        "4. List the most likely conditions in order of probability."
    ),
    expected_output="A detailed preliminary diagnosis with ranked possible conditions.",
    agent=diagnostician
)

treatment_task = Task(
    description=(
        "1. Based on the diagnosis, create a comprehensive treatment plan.\n"
        f"2. Consider patient profile: Age: {age}, Gender: {gender}\n"
        f"3. Account for medical history: {medical_history}\n"
        "4. Provide detailed recommendations including:\n"
        "   - Medications and dosages\n"
        "   - Lifestyle modifications\n"
        "   - Follow-up care schedule\n"
        "   - Warning signs to watch for"
    ),
    expected_output="A comprehensive, personalized treatment plan.",
    agent=treatment_advisor
)

# Create Crew
crew = Crew(
    agents=[diagnostician, treatment_advisor],
    tasks=[diagnose_task, treatment_task],
    verbose=True
)

# Analysis Button
if st.button("Generate Analysis"):
    if not symptoms or not medical_history:
        st.error("Please provide both symptoms and medical history before generating analysis.")
    else:
        with tab2:
            with st.status("🔄 Processing..."):
                st.write("Analyzing patient data...")
                st.write("Generating diagnosis...")
                st.write("Creating treatment plan...")
                
                result = crew.kickoff(inputs={
                    "symptoms": symptoms,
                    "medical_history": medical_history
                })
                
            st.success("Analysis Complete!")
            
            # Display Results
            st.markdown("### 📋 Analysis Results")
            if isinstance(result, tuple):
                for item in result:
                    st.markdown(str(item))
                    st.markdown("---")  # Add a separator between items
            else:
                st.markdown(str(result))
            
            # Generate and offer download
            docx_file = generate_docx(result)
            download_link = get_download_link(docx_file, "medical_analysis_report.docx")
            st.markdown("### 📥 Download Report")
            st.markdown(download_link, unsafe_allow_html=True)
            
            # Additional recommendations
            st.markdown("### ⚡ Next Steps")
            st.info("""
                1. Review the generated report in detail
                2. Consider additional specialist consultations if needed
                3. Schedule necessary follow-up appointments
                4. Monitor patient progress and adjust treatment as needed
            """)